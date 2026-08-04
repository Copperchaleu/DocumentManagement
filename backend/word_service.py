"""将项目信息合并写入 Word 文档。"""

from __future__ import annotations

import os
import re
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor
from docx.text.run import Run


class WordFileLockedError(Exception):
    """目标 Word 正在被占用（如被 Microsoft Word 打开）导致无法写入。"""

    def __init__(self, path: Path | str, cause: Exception | None = None) -> None:
        self.path = Path(path)
        self.cause = cause
        super().__init__(str(self.path))


class WordFileWriteError(Exception):
    """Word 写入失败（非占用类错误）。"""

    def __init__(self, path: Path | str, cause: Exception | None = None) -> None:
        self.path = Path(path)
        self.cause = cause
        super().__init__(str(self.path))


def is_file_locked(path: Path) -> bool:
    """
    检测文件是否被占用。
    Word/WPS 打开 docx 时，通常会拒绝覆盖/重命名（PermissionError / WinError 32）。
    """
    path = Path(path)
    if not path.exists():
        return False

    # 1) 尝试以读写方式打开
    try:
        with open(path, "r+b"):
            pass
    except PermissionError:
        return True
    except OSError as e:
        if getattr(e, "winerror", None) == 32 or e.errno in (13, 11):
            return True

    # 2) 重命名探测：被 Word 占用时往往无法改名
    probe = path.with_name(path.name + ".__lockprobe__")
    try:
        if probe.exists():
            try:
                probe.unlink()
            except Exception:
                pass
        path.rename(probe)
        probe.rename(path)
        return False
    except PermissionError:
        return True
    except OSError as e:
        if getattr(e, "winerror", None) == 32 or e.errno in (13, 11):
            return True
        # 其他异常保守视为可能占用，避免误覆盖
        return True
    finally:
        # 若 probe 残留且原文件不在，尝试回滚
        try:
            if probe.exists() and not path.exists():
                probe.rename(path)
        except Exception:
            pass


def assert_word_writable(path: Path) -> None:
    """写入前检查；若被占用则抛 WordFileLockedError。"""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists() and is_file_locked(path):
        raise WordFileLockedError(path)


def document_to_docx_bytes(doc: Document) -> bytes:
    """把 python-docx 文档序列化为可写入数据库的完整 docx 字节。"""
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def write_docx_bytes_atomic(file_content: bytes, output_path: Path) -> Path:
    """将 docx 字节先写临时文件再替换；占用时给出明确异常。"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    assert_word_writable(output_path)

    tmp_path = output_path.with_name(output_path.name + ".writing.tmp")
    try:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except Exception:
                pass
        tmp_path.write_bytes(file_content)
        try:
            os.replace(str(tmp_path), str(output_path))
        except PermissionError as e:
            raise WordFileLockedError(output_path, e) from e
        except OSError as e:
            if getattr(e, "winerror", None) == 32 or e.errno in (13, 11):
                raise WordFileLockedError(output_path, e) from e
            raise WordFileWriteError(output_path, e) from e
    except WordFileLockedError:
        raise
    except WordFileWriteError:
        raise
    except PermissionError as e:
        raise WordFileLockedError(output_path, e) from e
    except Exception as e:
        raise WordFileWriteError(output_path, e) from e
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except Exception:
                pass
    return output_path


def _atomic_save_docx(doc: Document, output_path: Path) -> Path:
    """兼容旧调用：先序列化文档，再原子写入。"""
    try:
        file_content = document_to_docx_bytes(doc)
    except Exception as e:
        raise WordFileWriteError(output_path, e) from e
    return write_docx_bytes_atomic(file_content, output_path)


def _set_run_font(run, size_pt: float | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _setup_document_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(6)


class _RichTextParser(HTMLParser):
    """将受控 wangEditor HTML 解析为段落块和带格式的文字片段。"""

    BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "li", "pre"}
    INLINE_TAGS = {"strong", "b", "em", "i", "u", "s", "strike", "del", "a", "span", "font"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[dict[str, Any]] = []
        self._block: dict[str, Any] | None = None
        self._style_stack: list[dict[str, Any]] = [self._base_style()]
        self._inline_tags: list[str] = []
        self._list_stack: list[str] = []

    @staticmethod
    def _base_style() -> dict[str, Any]:
        return {
            "bold": False,
            "italic": False,
            "underline": False,
            "strike": False,
            "color": None,
            "href": None,
        }

    @staticmethod
    def _attrs(attrs) -> dict[str, str]:
        return {str(k).lower(): str(v or "") for k, v in attrs}

    @staticmethod
    def _css(style_text: str) -> dict[str, str]:
        result: dict[str, str] = {}
        for part in (style_text or "").split(";"):
            if ":" not in part:
                continue
            key, value = part.split(":", 1)
            result[key.strip().lower()] = value.strip()
        return result

    @staticmethod
    def _normalize_color(value: str | None) -> str | None:
        raw = (value or "").strip().lower()
        if re.fullmatch(r"#[0-9a-f]{6}", raw):
            return raw[1:].upper()
        if re.fullmatch(r"#[0-9a-f]{3}", raw):
            return "".join(ch * 2 for ch in raw[1:]).upper()
        match = re.fullmatch(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", raw)
        if match:
            nums = [max(0, min(255, int(v))) for v in match.groups()]
            return "".join(f"{v:02X}" for v in nums)
        return None

    @staticmethod
    def _alignment(attrs: dict[str, str]) -> str:
        css = _RichTextParser._css(attrs.get("style", ""))
        align = (css.get("text-align") or attrs.get("align") or "left").lower()
        return align if align in {"left", "center", "right", "justify"} else "left"

    def _start_block(self, tag: str, attrs: dict[str, str]) -> None:
        self._flush_block()
        block_type = "list-item" if tag == "li" else tag
        self._block = {
            "type": block_type,
            "runs": [],
            "align": self._alignment(attrs),
            "list_type": self._list_stack[-1] if tag == "li" and self._list_stack else None,
            "list_depth": max(0, len(self._list_stack) - 1),
        }

    def _ensure_block(self) -> None:
        if self._block is None:
            self._start_block("p", {})

    def _flush_block(self) -> None:
        if self._block is None:
            return
        runs = self._block["runs"]
        while runs and not runs[-1]["text"]:
            runs.pop()
        if any(run["text"].strip() for run in runs):
            self.blocks.append(self._block)
        self._block = None

    def _append_text(self, text: str) -> None:
        clean = text.replace("\u00a0", " ")
        if not clean:
            return
        self._ensure_block()
        style = dict(self._style_stack[-1])
        runs = self._block["runs"]
        if runs and all(runs[-1].get(k) == style.get(k) for k in style):
            runs[-1]["text"] += clean
        else:
            runs.append({"text": clean, **style})

    def _push_inline_style(self, tag: str, attrs: dict[str, str]) -> None:
        style = dict(self._style_stack[-1])
        css = self._css(attrs.get("style", ""))
        if tag in {"strong", "b"} or css.get("font-weight", "").lower() in {"bold", "600", "700", "800", "900"}:
            style["bold"] = True
        if tag in {"em", "i"} or css.get("font-style", "").lower() == "italic":
            style["italic"] = True
        decoration = css.get("text-decoration", "").lower()
        if tag == "u" or "underline" in decoration:
            style["underline"] = True
        if tag in {"s", "strike", "del"} or "line-through" in decoration:
            style["strike"] = True
        color = self._normalize_color(css.get("color") or attrs.get("color"))
        if color:
            style["color"] = color
        if tag == "a" and attrs.get("href"):
            style["href"] = attrs["href"].strip()
        self._style_stack.append(style)
        self._inline_tags.append(tag)

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        attr_map = self._attrs(attrs)
        if tag in {"ul", "ol"}:
            self._list_stack.append(tag)
        elif tag in self.BLOCK_TAGS:
            self._start_block(tag, attr_map)
        elif tag in self.INLINE_TAGS:
            self._push_inline_style(tag, attr_map)
        elif tag == "br":
            self._append_text("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.BLOCK_TAGS:
            self._flush_block()
        elif tag in {"ul", "ol"}:
            if self._list_stack:
                self._list_stack.pop()
        elif tag in self.INLINE_TAGS and len(self._style_stack) > 1:
            if self._inline_tags:
                self._inline_tags.pop()
            self._style_stack.pop()

    def handle_data(self, data: str) -> None:
        self._append_text(data)

    def close(self) -> None:
        super().close()
        self._flush_block()


def _parse_rich_text(content: str) -> list[dict[str, Any]]:
    parser = _RichTextParser()
    parser.feed(content or "")
    parser.close()
    return parser.blocks


def html_to_plain_text(content: str) -> str:
    """富文本转纯文本，兼容历史纯文本数据。"""
    raw = content or ""
    if not re.search(r"<\s*[a-zA-Z][^>]*>", raw):
        return raw
    blocks = _parse_rich_text(raw)
    return "\n".join("".join(run["text"] for run in block["runs"]).strip() for block in blocks)


def _apply_run_style(run: Run, spec: dict[str, Any]) -> None:
    _set_run_font(run)
    run.bold = bool(spec.get("bold"))
    run.italic = bool(spec.get("italic"))
    run.underline = bool(spec.get("underline"))
    run.font.strike = bool(spec.get("strike"))
    color = spec.get("color")
    if color and re.fullmatch(r"[0-9A-F]{6}", color):
        run.font.color.rgb = RGBColor.from_string(color)


def _set_run_text(run: Run, text: str) -> None:
    parts = text.split("\n")
    for index, part in enumerate(parts):
        if index:
            run.add_break()
        if part:
            run.add_text(part)


def _add_hyperlink(paragraph, text: str, url: str, spec: dict[str, Any]) -> Run:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run_element = OxmlElement("w:r")
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    run = Run(run_element, paragraph)
    _set_run_text(run, text)
    link_spec = dict(spec)
    link_spec["color"] = link_spec.get("color") or "0563C1"
    link_spec["underline"] = True
    _apply_run_style(run, link_spec)
    return run


def _paragraph_alignment(value: str):
    return {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(value, WD_ALIGN_PARAGRAPH.LEFT)


def _add_block_runs(paragraph, runs: list[dict[str, Any]]) -> None:
    for spec in runs:
        text = spec.get("text") or ""
        if not text:
            continue
        href = (spec.get("href") or "").strip()
        if href:
            _add_hyperlink(paragraph, text, href, spec)
        else:
            run = paragraph.add_run()
            _set_run_text(run, text)
            _apply_run_style(run, spec)


def _add_rich_content(doc: Document, content: str) -> None:
    raw = content or ""
    if not re.search(r"<\s*[a-zA-Z][^>]*>", raw):
        _add_paragraphs(doc, raw)
        return

    for block in _parse_rich_text(raw):
        block_type = block["type"]
        if block_type in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = min(4, max(2, int(block_type[1]) + 1))
            paragraph = doc.add_heading(level=level)
        elif block_type == "list-item":
            base = "List Number" if block.get("list_type") == "ol" else "List Bullet"
            depth = min(2, int(block.get("list_depth") or 0))
            style = base if depth == 0 else f"{base} {depth + 1}"
            try:
                paragraph = doc.add_paragraph(style=style)
            except KeyError:
                paragraph = doc.add_paragraph(style=base)
        else:
            paragraph = doc.add_paragraph()

        paragraph.alignment = _paragraph_alignment(block.get("align") or "left")
        if block_type == "blockquote":
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.right_indent = Pt(8)
        _add_block_runs(paragraph, block["runs"])


def _add_paragraphs(doc: Document, content: str) -> None:
    lines = (content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        text = "\n".join(buffer).strip("\n")
        if text.strip() == "" and buffer:
            doc.add_paragraph("")
        elif text:
            p = doc.add_paragraph(text)
            for run in p.runs:
                _set_run_font(run)
        buffer = []

    for line in lines:
        if line.strip() == "":
            flush()
            buffer = [""]
            flush()
        else:
            buffer.append(line)
    flush()


def build_period_word_document(
    category_path_names: Iterable[str],
    period_type: str,
    period_label: str,
    projects: list[dict[str, Any]],
) -> bytes:
    """
    把同一分类 + 同一时间周期的所有项目生成成 docx 字节。
    返回值可先保存进 SQLite，再同步为本地最新文件。
    """
    period_cn = {"week": "周", "month": "月", "quarter": "季度"}.get(period_type, period_type)
    cat_path = " / ".join([n for n in category_path_names if n]) or "未分类"
    title = f"{cat_path} · {period_cn}报 {period_label}"

    doc = Document()
    _setup_document_style(doc)

    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        _set_run_font(run, color=RGBColor(0x1A, 0x1A, 0x2E))

    meta = doc.add_paragraph(
        f"分类：{cat_path}  |  周期：{period_cn} {period_label}  |  项目数：{len(projects)}"
    )
    for run in meta.runs:
        _set_run_font(run, size_pt=9, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph("")

    if not projects:
        p = doc.add_paragraph("（本周期暂无项目）")
        for run in p.runs:
            _set_run_font(run, color=RGBColor(0x88, 0x88, 0x88))
    else:
        for idx, project in enumerate(projects, start=1):
            # 项目标题
            h = doc.add_heading(f"{idx}. {project.get('title') or '未命名项目'}", level=2)
            for run in h.runs:
                _set_run_font(run)

            # 元信息
            bits = []
            if project.get("created_at"):
                bits.append(f"创建：{project['created_at']}")
            if project.get("updated_at"):
                bits.append(f"更新：{project['updated_at']}")
            atts = project.get("attachments") or []
            if atts:
                names = "、".join(a.get("original_name") or a.get("stored_name") or "" for a in atts)
                bits.append(f"附件：{names}")
            if bits:
                info = doc.add_paragraph("  |  ".join(bits))
                for run in info.runs:
                    _set_run_font(run, size_pt=9, color=RGBColor(0x66, 0x66, 0x66))

            _add_rich_content(doc, project.get("content") or "")
            doc.add_paragraph("")  # 项目之间留空

    return document_to_docx_bytes(doc)


def create_period_word_document(
    output_path: Path,
    category_path_names: Iterable[str],
    period_type: str,
    period_label: str,
    projects: list[dict[str, Any]],
) -> Path:
    """兼容旧接口：生成 docx 字节并写入本地文件。"""
    file_content = build_period_word_document(
        category_path_names=category_path_names,
        period_type=period_type,
        period_label=period_label,
        projects=projects,
    )
    return write_docx_bytes_atomic(file_content, output_path)


# 兼容旧接口（若被引用）
def create_word_document(
    title: str,
    content: str,
    output_path: Path,
    category_name: str = "",
    created_at: str = "",
) -> Path:
    project = {
        "title": title,
        "content": content,
        "created_at": created_at,
        "updated_at": created_at,
        "attachments": [],
    }
    return create_period_word_document(
        output_path=output_path,
        category_path_names=[category_name] if category_name else [],
        period_type="month",
        period_label="legacy",
        projects=[project],
    )
