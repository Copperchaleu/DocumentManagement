"""独立验证脚本：验证字体导出改动（后端）与前端工具栏白名单（静态）。

运行方式：
    cd /Users/graypaul/Projects/DocumentManagement && .venv/bin/python scripts/verify_font_export.py

说明：项目后端 venv 未安装 pytest，故使用原生 assert + 打印 PASS/FAIL 的方案，
任一断言失败即视为整体不通过（IS_PASS: NO）。
"""

from __future__ import annotations

import re
import sys

sys.path.insert(0, "backend")

import word_service  # noqa: E402
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import RGBColor  # noqa: E402

# --------------------------------------------------------------------------- #
# 运行统计
# --------------------------------------------------------------------------- #
_RESULTS: list[tuple[str, bool, str]] = []


def check(name: str, condition: bool, detail: str = "") -> None:
    _RESULTS.append((name, condition, detail))
    status = "PASS" if condition else "FAIL"
    print(f"[{status}] {name}" + (f"  -> {detail}" if detail else ""))


def approx(a: float, b: float, tol: float = 0.01) -> bool:
    return abs(a - b) <= tol


# --------------------------------------------------------------------------- #
# 工具：读取 run 的 rFonts/eastAsia / 字号(pt)
# --------------------------------------------------------------------------- #
def _rfonts(run):
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return None
    return rPr.find(qn("w:rFonts"))


def _east_asia(run):
    rf = _rfonts(run)
    if rf is None:
        return None
    return rf.get(qn("w:eastAsia"))


def _run_pt(run):
    if run.font.size is None:
        return None
    return run.font.size.pt


def _find_run_by_text(doc: Document, needle: str):
    for para in doc.paragraphs:
        for run in para.runs:
            if needle in (run.text or ""):
                return run
    return None


def _all_runs(doc: Document):
    runs = []
    for para in doc.paragraphs:
        runs.extend(para.runs)
    return runs


# =========================================================================== #
# A) 后端验证（必须实跑）
# =========================================================================== #

# --- A1: 中文字体 + px -------------------------------------------------------
def test_font_cjk_px():
    blocks = word_service._parse_rich_text(
        '<span style="font-family: 微软雅黑; font-size: 16px;">文本</span>'
    )
    found = [r for b in blocks for r in b["runs"] if r["text"] == "文本"]
    ok = bool(found) and found[0]["font_family"] == "微软雅黑" and approx(found[0]["font_size"], 12.0)
    detail = f"font_family={found[0]['font_family'] if found else None}, font_size={found[0]['font_size'] if found else None}"
    check("A1 中文字体+px 解析", ok, detail)


# --- A2: 拉丁字体 + px -------------------------------------------------------
def test_font_latin_px():
    blocks = word_service._parse_rich_text(
        '<span style="font-family: Arial; font-size: 24px;">x</span>'
    )
    found = [r for b in blocks for r in b["runs"] if r["text"] == "x"]
    ok = bool(found) and found[0]["font_family"] == "Arial" and approx(found[0]["font_size"], 18.0)
    detail = f"font_family={found[0]['font_family'] if found else None}, font_size={found[0]['font_size'] if found else None}"
    check("A2 拉丁字体+px 解析", ok, detail)


# --- A3: 带引号字体名 + px ---------------------------------------------------
def test_font_quoted_px():
    blocks = word_service._parse_rich_text(
        "<span style=\"font-family: 'Times New Roman'; font-size: 20px;\">Hello</span>"
    )
    found = [r for b in blocks for r in b["runs"] if r["text"] == "Hello"]
    ok = bool(found) and found[0]["font_family"] == "Times New Roman" and approx(found[0]["font_size"], 15.0)
    detail = f"font_family={found[0]['font_family'] if found else None}, font_size={found[0]['font_size'] if found else None}"
    check("A3 引号字体名+px 解析", ok, detail)


# --- A4: 端到端落盘 ----------------------------------------------------------
def test_end_to_end_docx():
    content = (
        '<span style="font-family: 微软雅黑; font-size: 16px;">中文</span>'
        '<span style="font-family: Arial; font-size: 24px;">EN</span>'
    )
    data = word_service.build_period_word_document(
        category_path_names=["测试分类"],
        period_type="month",
        period_label="2026-08",
        projects=[
            {
                "title": "T",
                "content": content,
                "created_at": "2026-08-04",
                "updated_at": "2026-08-04",
                "attachments": [],
            }
        ],
    )
    doc = Document(__import__("io").BytesIO(data))

    cn_run = _find_run_by_text(doc, "中文")
    en_run = _find_run_by_text(doc, "EN")

    cn_ok = (
        cn_run is not None
        and (cn_run.font.name == "微软雅黑" or (cn_run.font.name or "").endswith("微软雅黑"))
        and _east_asia(cn_run) == "微软雅黑"
        and approx(_run_pt(cn_run), 12.0)
    )
    en_ok = (
        en_run is not None
        and en_run.font.name == "Arial"
        and _east_asia(en_run) == "微软雅黑"
        and approx(_run_pt(en_run), 18.0)
    )
    detail = (
        f"中文 run.name={getattr(cn_run,'font',None) and cn_run.font.name}, "
        f"eastAsia={_east_asia(cn_run)}, pt={_run_pt(cn_run)} | "
        f"EN run.name={en_run.font.name if en_run else None}, "
        f"eastAsia={_east_asia(en_run)}, pt={_run_pt(en_run)}"
    )
    check("A4 端到端落盘（中文/EN 字体字号 eastAsia）", cn_ok and en_ok, detail)


# --- A5: 回归——旧特性仍在 ---------------------------------------------------
def test_regression_old_features():
    html = (
        "<b>粗体</b><i>斜体</i><u>下划线</u><s>删除</s>"
        '<span style="color: #ff0000;">红字</span>'
        '<a href="http://example.com">链接</a>'
    )
    data = word_service.build_period_word_document(
        category_path_names=["测试"],
        period_type="month",
        period_label="2026-08",
        projects=[
            {
                "title": "R",
                "content": html,
                "created_at": "2026-08-04",
                "updated_at": "2026-08-04",
                "attachments": [],
            }
        ],
    )
    doc = Document(__import__("io").BytesIO(data))
    runs = _all_runs(doc)

    bold = any(getattr(r, "bold", False) for r in runs)
    italic = any(getattr(r, "italic", False) for r in runs)
    underline = any(getattr(r, "underline", False) for r in runs)
    strike = any(getattr(r, "font", None) and getattr(r.font, "strike", False) for r in runs)
    red = any(
        getattr(r.font, "color", None) is not None and r.font.color.rgb == RGBColor(0xFF, 0, 0)
        for r in runs
    )

    # 超链接关系：document.part.rels 中存在 external hyperlink
    hyperlink = any(
        rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
        and rel.is_external
        and str(rel.target_ref) == "http://example.com"
        for rel in doc.part.rels.values()
    )

    ok = bold and italic and underline and strike and red and hyperlink
    detail = (
        f"bold={bold}, italic={italic}, underline={underline}, "
        f"strike={strike}, red={red}, hyperlink={hyperlink}"
    )
    check("A5 回归（粗体/斜体/下划线/删除/红色/超链接）", ok, detail)


# --- A6: 不支持的 HTML 不破坏解析 -------------------------------------------
def test_unsupported_html():
    content = (
        '<img src="x.png">视频<video src="y.mp4"></video>'
        "测试😀表情"
        "<iframe src='http://evil.com'></iframe>"
    )
    # 1) _parse_rich_text 不抛异常
    try:
        blocks = word_service._parse_rich_text(content)
        parsed_text = "".join(r["text"] for b in blocks for r in b["runs"])
        parse_ok = "测试表情" in parsed_text or "测试" in parsed_text
    except Exception as exc:  # noqa: BLE001
        blocks, parsed_text, parse_ok = None, "", False
        detail_parse = f"EXCEPTION: {exc!r}"
        check("A6 _parse_rich_text 容错", False, detail_parse)
        return

    # 2) build_period_word_document 不抛异常
    try:
        data = word_service.build_period_word_document(
            category_path_names=["测试"],
            period_type="month",
            period_label="2026-08",
            projects=[
                {
                    "title": "U",
                    "content": content,
                    "created_at": "2026-08-04",
                    "updated_at": "2026-08-04",
                    "attachments": [],
                }
            ],
        )
        doc = Document(__import__("io").BytesIO(data))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        build_ok = "测试表情" in full_text or "表情" in full_text
    except Exception as exc:  # noqa: BLE001
        data, full_text, build_ok = None, "", False
        detail_build = f"EXCEPTION: {exc!r}"
        check("A6 build_period_word_document 容错", False, detail_build)
        return

    ok = parse_ok and build_ok
    check("A6 不支持的 HTML 不破坏解析（img/video/emoji/iframe）", ok,
          f"解析含'测试表情'={parse_ok}, 落盘含'表情'={build_ok}")


# --- A7: _is_cjk_font --------------------------------------------------------
def test_is_cjk_font():
    cases = {
        "微软雅黑": True,
        "SimSun": True,
        "宋体": True,
        "Arial": False,
        "Times New Roman": False,
    }
    ok = all(word_service._is_cjk_font(name) is expected for name, expected in cases.items())
    detail = ", ".join(f"{k}->{word_service._is_cjk_font(k)}" for k in cases)
    check("A7 _is_cjk_font 判定", ok, detail)


# =========================================================================== #
# B) 前端静态校验（读文件 + 字符串断言，不执行 npm build）
# =========================================================================== #
def test_frontend_toolbar():
    vue_path = "web/src/views/ComposeView.vue"
    try:
        with open(vue_path, encoding="utf-8") as fh:
            src = fh.read()
    except Exception as exc:  # noqa: BLE001
        check("B 前端文件读取", False, f"EXCEPTION: {exc!r}")
        return

    # 1) editorMode 为 'default'
    m = re.search(r"const\s+editorMode\s*=\s*'([^']+)'", src)
    editor_mode = m.group(1) if m else None
    check("B1 editorMode=='default'", editor_mode == "default", f"editorMode={editor_mode}")

    # 2) toolbarKeys 包含 fontFamily 与 fontSize
    has_font_family = "'fontFamily'" in src and re.search(r"toolbarKeys", src) is not None
    # 更稳妥：提取 toolbarKeys 数组内容
    mkeys = re.search(r"toolbarKeys\s*:\s*\[(.*?)\]", src, re.DOTALL)
    if mkeys:
        keys_block = mkeys.group(1)
        keys = re.findall(r"'([^']+)'", keys_block)
        has_family = "fontFamily" in keys
        has_size = "fontSize" in keys
        detail = f"keys={keys}"
    else:
        has_family = has_size = False
        detail = "toolbarKeys 未找到"
    check("B2 toolbarKeys 含 fontFamily/fontSize", has_family and has_size, detail)

    # 3) toolbarKeys 不包含禁用键
    forbidden = [
        "insertImage", "insertVideo", "bgColor", "insertTable", "emoticon",
        "todo", "codeBlock", "indent", "delIndent", "fullScreen", "divider",
        "group-image", "group-video",
    ]
    present_forbidden = [k for k in forbidden if (mkeys and k in keys)]
    check("B3 toolbarKeys 不含禁用键", len(present_forbidden) == 0,
          f"命中禁用键={present_forbidden}")

    # 4) 提示文字
    hint_ok = ("字体" in src and "字号" in src
               and ("图片" in src and "视频" in src and "表情" in src)
               and "支持标题、粗体、列表、引用和链接等常用格式" not in src)
    check("B4 提示文字（字体/字号 + 图片/视频/表情已移除 + 旧文案不存在）", hint_ok,
          "")


# =========================================================================== #
# 主流程
# =========================================================================== #
def main() -> int:
    print("=" * 70)
    print("后端字体导出验证")
    print("=" * 70)
    test_font_cjk_px()
    test_font_latin_px()
    test_font_quoted_px()
    test_end_to_end_docx()
    test_regression_old_features()
    test_unsupported_html()
    test_is_cjk_font()

    print("-" * 70)
    print("前端工具栏静态校验")
    print("-" * 70)
    test_frontend_toolbar()

    print("=" * 70)
    total = len(_RESULTS)
    passed = sum(1 for _, ok, _ in _RESULTS if ok)
    failed = total - passed
    print(f"汇总: 总 {total} | 通过 {passed} | 失败 {failed}")
    for name, ok, detail in _RESULTS:
        if not ok:
            print(f"  失败项: {name}  {detail}")

    is_pass = failed == 0
    print(f"IS_PASS: {'YES' if is_pass else 'NO'}")
    print("=" * 70)
    return 0 if is_pass else 1


if __name__ == "__main__":
    sys.exit(main())
